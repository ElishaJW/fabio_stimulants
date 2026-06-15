"""Generate diagnostic report as .docx"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Style configuration ---
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

for level, (size, space_before) in enumerate([(18, 24), (14, 18), (12, 14)], 1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Arial'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    h.font.bold = True
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(8)

# ── TITLE PAGE ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run('Diagnostic Report')
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
run.font.bold = True
run.font.name = 'Arial'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Discrepancies Between National and Ecoregion-Scale\nBiodiversity Impact Estimates for Stimulant Commodities')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
run.font.name = 'Arial'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('Prepared by Claude (analytical assistant) for Eli Wilson\nApril 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Arial'

doc.add_page_break()

# ── HELPER FUNCTIONS ──
def add_body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(11)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_severity_tag(text, color_hex):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = 'Arial'
    r.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    return p

def add_recommendation(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        pass
    else:
        p.clear()
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        r2.font.size = Pt(11)
    return p

# ── SECTION 1: EXECUTIVE SUMMARY ──
doc.add_heading('1. Executive Summary', level=1)

add_body(
    'Three anomalies have been identified when comparing national-scale and ecoregion-scale '
    'biodiversity impact results for stimulant commodity production (coffee, tea, cocoa, tobacco):'
)

items = [
    'At the national scale, tea\u2019s biodiversity impact is dominated by freshwater eutrophication '
    '(N and P application) attributed primarily to Sri Lanka. Land use is the dominant stressor '
    'overall when summed across all commodities.',

    'At the ecoregion scale, Sri Lanka\u2019s most impacted ecoregion (\u201cSri Lanka dry-zone dry '
    'evergreen forests\u201d) ranks only 4th among the top 20 most impacted ecoregions globally \u2014 '
    'despite Sri Lanka being the highest-impact consuming country at the national scale.',

    'At the ecoregion scale, eutrophication dominates the biodiversity impact in virtually all '
    'top 20 ecoregions, contradicting the national-scale finding that land use is the dominant '
    'stressor globally.'
]
for item in items:
    p = doc.add_paragraph(item, style='List Number')
    for r in p.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(11)

add_body(
    'This report identifies six root causes ranging from a critical methodological error in '
    'characterization factor selection to structural biases in the realm-averaging formula. '
    'Findings are ranked by estimated severity.'
)

# ── SECTION 2: ROOT CAUSE ANALYSIS ──
doc.add_heading('2. Root Cause Analysis', level=1)

# --- 2.1 ---
doc.add_heading('2.1  Use of \u201cDirect\u201d vs. \u201cDiffuse\u201d Eutrophication Characterization Factors', level=2)
add_severity_tag('Severity: HIGH \u2014 Most likely primary cause of all three anomalies', 'C0392B')

add_body(
    'The country-level biodiversity calculations (17_biodiv_calculations.R, lines 276\u2013278) '
    'apply country-averaged eutrophication characterization factors (CFs) from the Helmes et al. '
    'freshwater eutrophication CF dataset. Specifically, the column used is \u201cCF_avg_diff\u201d \u2014 '
    'the average CF for diffuse sources. Diffuse-source CFs model nutrient transport via soil '
    'leaching and surface runoff, which is the appropriate emission pathway for agricultural '
    'fertilizer application.'
)

add_body(
    'The ecoregion-level calculations (plot_fig_3_bd_footprint_ecoregion_calc.R, lines 112\u2013113) '
    'load spatially resolved raster CFs from files named CF_average_direct.asc for both N and P. '
    'These are direct-source CFs, which model nutrients discharged directly into waterways '
    '(e.g., point sources such as wastewater treatment plants or direct industrial discharge).'
)

add_body(
    'Direct-source CFs are systematically higher than diffuse-source CFs because they do not '
    'account for landscape attenuation \u2014 the physical, chemical, and biological processes that '
    'reduce nutrient loads as they travel through soil and landscape before reaching surface water. '
    'For agricultural N and P application from stimulant crop fertilization, the diffuse pathway '
    'is the correct emission model. Using direct-source CFs overstates the fraction of applied '
    'nutrients that reaches freshwater ecosystems, thereby overestimating eutrophication '
    'biodiversity impacts relative to other stressors (particularly land use).'
)

add_body(
    'The correct diffuse-source rasters (CF_average_diffuse.asc) exist in the same directory '
    'for both N and P. This was confirmed by listing the contents of:', bold_prefix='Confirmation: '
)
add_code('X:/Eli/DATA/cf/freshwater eutrophication/.../N/ASCII_rasters/')
add_code('X:/Eli/DATA/cf/freshwater eutrophication/.../P/ASCII_rasters/')

add_body(
    'Replace CF_average_direct.asc with CF_average_diffuse.asc in '
    'plot_fig_3_bd_footprint_ecoregion_calc.R (lines 112\u2013113). This single change is expected '
    'to substantially reduce the eutrophication dominance across all top 20 ecoregions.',
    bold_prefix='Recommendation: '
)

# --- 2.2 ---
doc.add_heading('2.2  Structural Bias in Realm Averaging Formula', level=2)
add_severity_tag('Severity: MEDIUM', 'E8A838')

add_body(
    'The combined biodiversity footprint at ecoregion scale is computed as '
    '(plot_fig_3_bd_footprint_ecoregion_calc.R, line 620):'
)
add_code('combined_bd = (landuse_bd + fw_total_eco) / 2')
add_code('where fw_total_eco = water_bd + n_bd + p_bd')

add_body(
    'The freshwater realm aggregates three distinct stressors (blue water consumption, nitrogen '
    'eutrophication, phosphorus eutrophication) into a single sum, while the terrestrial realm '
    'contains only one stressor (land use). This creates an asymmetry: even if each individual '
    'freshwater stressor is smaller in magnitude than land use, their sum can exceed land use '
    'by a factor of 2\u20133\u00d7, biasing the combined score toward freshwater dominance.'
)

add_body(
    'This asymmetry also creates a selection bias in the top-20 ecoregion ranking. Ecoregions '
    'that rank highest by combined_bd are disproportionately those where freshwater impacts '
    'happen to be anomalously large \u2014 making eutrophication dominance a partially self-fulfilling '
    'artifact of the ranking methodology.'
)

add_body(
    'Note that the country-level scripts (plot_fig_1, plot_fig_2) include GHG impacts in the '
    'terrestrial realm:'
)
add_code('bd_fp_terrestrial = bd_fp_landuse + bd_fp_GHG')
add_code('bd_fp_total = (bd_fp_fw + bd_fp_terrestrial) / 2')

add_body(
    'This partially balances the realm accounting at the national scale but is absent at '
    'ecoregion scale (see Finding 2.4).'
)

add_body('Recommendations:', bold_prefix='')
recs = [
    'Include GHG impacts in the terrestrial realm at ecoregion scale to match the country-level methodology.',
    'Consider whether the three freshwater stressors should be averaged (rather than summed) before combining with terrestrial.',
    'At minimum, present decomposed results (landuse_bd and fw_total_eco separately) alongside combined_bd.'
]
for rec in recs:
    p = doc.add_paragraph(rec, style='List Bullet')
    for r in p.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(11)

# --- 2.3 ---
doc.add_heading('2.3  Sri Lanka\u2019s Extreme Country-Level N Characterization Factor', level=2)
add_severity_tag('Severity: MEDIUM', 'E8A838')

add_body(
    'Sri Lanka\u2019s country-level nitrogen CF for diffuse agricultural sources is approximately '
    '575 times higher than India\u2019s:'
)

# CF comparison table
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Country', 'N CF (PDF\u00b7yr/kg)', 'Relative to India']
data = [
    ['Sri Lanka', '3.27 \u00d7 10\u207b\u00b9\u00b9', '575\u00d7'],
    ['India', '5.69 \u00d7 10\u207b\u00b9\u2074', 'baseline'],
    ['China', '3.44 \u00d7 10\u207b\u00b9\u2074', '0.6\u00d7'],
    ['Kenya', '2.49 \u00d7 10\u207b\u00b9\u00b3', '4.4\u00d7'],
    ['Vietnam', '3.85 \u00d7 10\u207b\u00b9\u2074', '0.7\u00d7'],
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for r in cell.paragraphs[0].runs:
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(10)
for row_i, row_data in enumerate(data):
    for col_i, val in enumerate(row_data):
        cell = table.rows[row_i + 1].cells[col_i]
        cell.text = val
        for r in cell.paragraphs[0].runs:
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            if row_i == 0:
                r.bold = True

doc.add_paragraph()  # spacer

add_body(
    'This extreme value, combined with Sri Lanka\u2019s high N application intensity for tea '
    '(225 kg N/ha vs. 125 for India and 15 for Kenya), drives the national-scale finding '
    'that tea/Sri Lanka dominates eutrophication impacts.'
)

add_body(
    'At ecoregion scale, the spatially resolved raster CFs do not exhibit this same extreme '
    'value for Sri Lanka pixels. The country-level CF is an area-weighted aggregate that may be '
    'disproportionately influenced by a small number of highly sensitive freshwater catchments '
    'with concentrated endemic species. The raster CF provides finer spatial resolution and '
    'likely assigns lower CF values to the specific pixels where tea production occurs.'
)

add_body(
    'Additionally, Sri Lanka\u2019s country-level P CF is reported as NaN (missing) in the source '
    'data, meaning phosphorus eutrophication impacts from Sri Lanka are effectively zero at '
    'the national scale \u2014 another data quality concern.'
)

add_body(
    'Verify the Sri Lanka N CF against the original Helmes et al. (2012) or Azevedo et al. '
    '(2013) source data. Consider whether the country-level aggregate is appropriate for '
    'Sri Lanka\u2019s geography or whether the raster-based values provide a more reliable estimate.',
    bold_prefix='Recommendation: '
)

# --- 2.4 ---
doc.add_heading('2.4  Climate Change (GHG) Impacts Excluded from Ecoregion Analysis', level=2)
add_severity_tag('Severity: LOW\u2013MEDIUM', '4A90C4')

add_body(
    'The ecoregion-level analysis processes four stressors: land use, blue water, N application, '
    'and P application. Climate change biodiversity impacts (CH4_bd, CO2_bd, N2O_bd) are '
    'entirely absent.'
)

add_body(
    'The country-level scripts include GHG in the terrestrial realm. The absence of GHG from '
    'the ecoregion terrestrial side means the terrestrial realm is systematically underweighted, '
    'exacerbating the freshwater dominance described in Finding 2.2.'
)

add_body(
    'Spatially resolving climate change impacts to ecoregions is methodologically challenging '
    '(GHG impacts are global, not local), which likely explains the omission. However, the '
    'asymmetric treatment between scales should be acknowledged when comparing results.'
)

# --- 2.5 ---
doc.add_heading('2.5  CH4/CO2 Label Swap Bug in E_bd Matrix', level=2)
add_severity_tag('Severity: LOW \u2014 Does not affect ecoregion analysis', '5E813F')

add_body(
    'In 17_biodiv_calculations.R, lines 436\u2013437, the CH4 and CO2 biodiversity impact rows '
    'are constructed with swapped labels:'
)
add_code('bd_rows["CO2_bd", ] <- E_year["CH4", ] * cc_cf$CF["methane"]     # should be CH4_bd')
add_code('bd_rows["CH4_bd", ] <- E_year["CO2", ] * cc_cf$CF["CO2"]         # should be CO2_bd')

add_body(
    'Multiple downstream scripts note this bug with commented-out correction code '
    '(plot_fig_1 line 69, plot_fig_2 line 67). Since these rows are always summed together '
    'for GHG totals, the combined GHG impact is numerically correct \u2014 only the individual '
    'CH4 vs. CO2 attribution is swapped. This does not affect the ecoregion discrepancy '
    '(which excludes GHG entirely) but should be corrected for analytical transparency.'
)

# --- 2.6 ---
doc.add_heading('2.6  Methodological Consistency Between Scales', level=2)
add_severity_tag('Severity: LOW \u2014 Not a source of discrepancy', '5E813F')

add_body(
    'Both scales derive their stressor footprints from the same FABIO supply chain model '
    '(E matrix \u00d7 Leontief inverse \u00d7 final demand). The ecoregion analysis loads '
    'str_footprint_YEAR.rds files, which are pre-computed as E_diag \u00d7 L \u00d7 Y \u2014 the same '
    'mathematical operation used by the country-level scripts. Total footprints should '
    'therefore be consistent between scales.'
)

add_body(
    'The additional step in the ecoregion analysis is spatial disaggregation: country-level '
    'footprints are distributed across 10\u00d710 km production raster cells, then aggregated to '
    'ecoregion boundaries. This introduces spatial uncertainty dependent on the quality and '
    'resolution of the production rasters (rel_prod_normalized_*.tif). However, this step '
    'preserves totals and is not a source of the observed discrepancy.'
)

# ── SECTION 3: METHODOLOGICAL COMPARISON TABLE ──
doc.add_heading('3. Summary of Methodological Differences', level=1)

table_data = [
    ['Aspect', 'Country-Level Approach', 'Ecoregion-Level Approach', 'Impact on Discrepancy'],
    ['N/P Characterization Factors',
     'Diffuse-source CFs (Country_CF xlsx, "CF_avg_diff" column)',
     'Direct-source CFs (CF_average_direct.asc raster)',
     'Direct CFs are systematically higher, inflating eutrophication at ecoregion scale. PRIMARY DRIVER.'],
    ['CF Spatial Resolution',
     'Single value per country',
     'Raster at ~0.5\u00b0 resolution',
     'Raster CFs smooth out country-level extremes (e.g., Sri Lanka).'],
    ['Realm Averaging',
     '(FW + terrestrial) / 2, where terrestrial = land use + GHG',
     '(FW + terrestrial) / 2, where terrestrial = land use only',
     'Terrestrial underweighted at ecoregion scale (missing GHG).'],
    ['Freshwater Composition',
     'water + N + P (3 stressors summed)',
     'water + N + P (3 stressors summed)',
     'Consistent, but 3-vs-1 asymmetry amplifies FW dominance.'],
    ['GHG Impacts',
     'Included (CH4, CO2, N2O in terrestrial realm)',
     'Excluded',
     'Ecoregion terrestrial realm systematically lower.'],
    ['Supply Chain Model',
     'E_bd \u00d7 L \u00d7 Y (full Leontief)',
     'str_footprint (= E_diag \u00d7 L \u00d7 Y)',
     'Mathematically equivalent \u2014 not a source of discrepancy.'],
    ['Spatial Attribution',
     'By consuming country',
     'By production ecoregion (via 10 km rasters)',
     'Different units but totals preserved.'],
]

num_rows = len(table_data)
num_cols = 4
table2 = doc.add_table(rows=num_rows, cols=num_cols)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

for row_i, row_data in enumerate(table_data):
    for col_i, val in enumerate(row_data):
        cell = table2.rows[row_i].cells[col_i]
        cell.text = val
        for r in cell.paragraphs[0].runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            if row_i == 0:
                r.bold = True

# ── SECTION 4: RECOMMENDATIONS ──
doc.add_page_break()
doc.add_heading('4. Prioritized Recommendations', level=1)

recommendations = [
    ('IMMEDIATE FIX: ',
     'Replace CF_average_direct.asc with CF_average_diffuse.asc in '
     'plot_fig_3_bd_footprint_ecoregion_calc.R (lines 112\u2013113). This is the single '
     'highest-impact correction and aligns the ecoregion analysis with the emission pathway '
     'appropriate for agricultural nutrient application.'),
    ('SHORT-TERM: ',
     'Review the realm-averaging formula. Consider (a) including GHG in the ecoregion '
     'terrestrial realm, (b) reporting decomposed stressor contributions alongside the '
     'combined metric, and (c) evaluating whether summing vs. averaging freshwater '
     'sub-stressors is more defensible.'),
    ('SHORT-TERM: ',
     'Fix the CH4/CO2 label swap in 17_biodiv_calculations.R lines 436\u2013437. Uncomment '
     'the correction code in downstream scripts or fix at the source.'),
    ('MEDIUM-TERM: ',
     'Verify Sri Lanka\u2019s country-level N CF (3.27 \u00d7 10\u207b\u00b9\u00b9) against the '
     'original LCIA source. Consider whether this value is an outlier requiring correction '
     '(similar to the basin 21611 water CF correction already implemented) or a legitimate '
     'representation of Sri Lanka\u2019s freshwater biodiversity vulnerability.'),
    ('LONG-TERM: ',
     'Conduct a sensitivity analysis comparing results with direct vs. diffuse CFs, and with '
     'different realm-averaging formulations, to quantify the magnitude of these methodological '
     'choices on final rankings.'),
]

for i, (prefix, body) in enumerate(recommendations, 1):
    p = doc.add_paragraph(style='List Number')
    p.clear()
    r = p.add_run(prefix)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r = p.add_run(body)
    r.font.name = 'Arial'
    r.font.size = Pt(11)

# ── SECTION 5: CONCLUSION ──
doc.add_heading('5. Conclusion', level=1)

add_body(
    'The discrepancies between national and ecoregion-scale results are primarily driven by '
    'the use of direct-source (rather than diffuse-source) eutrophication characterization '
    'factors in the ecoregion analysis. This methodological mismatch inflates freshwater '
    'biodiversity impacts across all ecoregions. The effect is compounded by a structural '
    'asymmetry in the realm-averaging formula (3 freshwater stressors vs. 1 terrestrial '
    'stressor) and the exclusion of climate change from the ecoregion terrestrial realm.'
)

add_body(
    'Sri Lanka\u2019s anomalous rank change between scales is explained by the difference between '
    'its extreme country-level CF aggregate and the more spatially nuanced raster-based CFs.'
)

add_body(
    'Correcting the CF raster selection (direct \u2192 diffuse) is expected to be the single most '
    'impactful fix. The remaining structural issues (realm averaging, GHG inclusion) warrant '
    'methodological review but are secondary in magnitude.'
)

# Save
out_path = os.path.join(os.path.dirname(__file__),
                        'diagnostic_report_national_vs_ecoregion_BD.docx')
doc.save(out_path)
print(f'Report saved to: {out_path}')
